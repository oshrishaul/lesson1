# 1. Will the following code block cause an exception? If so, why?
# colors = ["red", "green", "blue", "purple"]
# for i in range(len(colors)):
#     print(colors[5])
# # ANSWER - In line 4 we are trying to print the fifth color in our list, which counts 4 colors. And this is why we are getting an "IndexError" exception.
#
# # 2. Suggest an edit to fix the above code.
# # ANSWER - we can use only 0 - 3 to print the required color or alternatively, add more colors to the string.
#
# # 3. What is the main risk with the code below:
my_file = open("C:/test/README.txt").read()
print(my_file)
# ANSWER - The main risk is the file will not be at the following location.
# 4. How would you protect the above code? And why?
try:
    my_file = open("C:/test/README.txt").read()
except FileNotFoundError:
    print("File not found")

# 5. Write a program which will:
# A. Create a folder named “test” (anywhere on your computer)
# B. Create a file named “name” containing your name inside.
import os
os.mkdir("C:/Users/shaulo2/Desktop/Test")
open("C:/Users/shaulo2/Desktop/Test/1.txt", "w+", encoding='utf-8').write("My name is Oshri Shaul")

# 6. A. Install python-docx python library through CMD /Terminal:
#    B. Create a docx file using the documentation: https://python-docx.readthedocs.io/en/latest/user/documents.html
from docx import Document
my_doc = Document()
my_doc.add_heading("Class 4 - Extra Assignment")
my_doc.add_paragraph("My name is Oshri Shaul")
my_doc.save("C:/Users/shaulo2/Desktop/Test/1.docx")


